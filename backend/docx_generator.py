"""
DOCX Generator - Creates perfectly formatted Word documents
Professional corporate resume formatting
"""

import io
import re
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy


def add_horizontal_line(paragraph):
    """Add a horizontal line below a paragraph"""
    p = paragraph._p
    pPr = p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '2C3E7A')
    pBdr.append(bottom)
    pPr.append(pBdr)


def set_paragraph_spacing(paragraph, before=0, after=0, line_spacing=None):
    """Set spacing for a paragraph"""
    pf = paragraph.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    if line_spacing:
        pf.line_spacing = Pt(line_spacing)


def add_run_with_style(paragraph, text, bold=False, italic=False,
                        font_name="Calibri", font_size=10, color=None):
    """Add a styled run to a paragraph"""
    run = paragraph.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.name = font_name
    run.font.size = Pt(font_size)
    if color:
        run.font.color.rgb = RGBColor(*color)
    return run


def generate_docx(resume_data: dict) -> bytes:
    """
    Generate a professionally formatted DOCX resume
    Returns bytes of the DOCX file
    """
    doc = Document()
    
    # ── Page Setup ──────────────────────────────────────────────────────────
    section = doc.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    
    # ── Color Constants ──────────────────────────────────────────────────────
    COLOR_HEADING = RGBColor(0x1A, 0x3A, 0x6B)   # Dark professional blue
    COLOR_SUBHEADING = RGBColor(0x2C, 0x2C, 0x2C) # Near black
    COLOR_BODY = RGBColor(0x33, 0x33, 0x33)        # Dark gray
    COLOR_META = RGBColor(0x55, 0x55, 0x55)        # Medium gray
    
    FONT_NAME = "Calibri"
    
    name = resume_data.get("name", "Your Name")
    contact = resume_data.get("contact", {})
    summary = resume_data.get("summary", "")
    skills = resume_data.get("skills", [])
    experience = resume_data.get("experience", [])
    education = resume_data.get("education", "")
    projects = resume_data.get("projects", "")
    certifications = resume_data.get("certifications", "")
    
    # ── NAME ─────────────────────────────────────────────────────────────────
    name_para = doc.add_paragraph()
    name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(name_para, before=0, after=4)
    name_run = name_para.add_run(name.upper() if name else "YOUR NAME")
    name_run.bold = True
    name_run.font.name = FONT_NAME
    name_run.font.size = Pt(20)
    name_run.font.color.rgb = COLOR_HEADING
    
    # ── CONTACT INFO ──────────────────────────────────────────────────────────
    contact_parts = []
    if contact.get("email"):
        contact_parts.append(contact["email"])
    if contact.get("phone"):
        contact_parts.append(contact["phone"])
    if contact.get("location"):
        contact_parts.append(contact["location"])
    if contact.get("linkedin"):
        contact_parts.append(contact["linkedin"])
    if contact.get("github"):
        contact_parts.append(contact["github"])
    
    if contact_parts:
        contact_para = doc.add_paragraph()
        contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_paragraph_spacing(contact_para, before=2, after=8)
        contact_run = contact_para.add_run(" | ".join(contact_parts))
        contact_run.font.name = FONT_NAME
        contact_run.font.size = Pt(9)
        contact_run.font.color.rgb = COLOR_META
    
    def add_section_header(doc, title):
        """Add a formatted section header with underline"""
        p = doc.add_paragraph()
        set_paragraph_spacing(p, before=10, after=3)
        run = p.add_run(title.upper())
        run.bold = True
        run.font.name = FONT_NAME
        run.font.size = Pt(11)
        run.font.color.rgb = COLOR_HEADING
        add_horizontal_line(p)
        return p
    
    # ── PROFESSIONAL SUMMARY ──────────────────────────────────────────────────
    if summary:
        add_section_header(doc, "Professional Summary")
        sum_para = doc.add_paragraph()
        set_paragraph_spacing(sum_para, before=3, after=3)
        sum_para.paragraph_format.first_line_indent = Inches(0)
        sum_run = sum_para.add_run(summary)
        sum_run.font.name = FONT_NAME
        sum_run.font.size = Pt(10)
        sum_run.font.color.rgb = COLOR_BODY
    
    # ── SKILLS ───────────────────────────────────────────────────────────────
    if skills:
        add_section_header(doc, "Core Competencies & Technical Skills")
        
        # Organize skills in rows of ~4 skills each
        skills_para = doc.add_paragraph()
        set_paragraph_spacing(skills_para, before=3, after=3)
        
        # Format as: "Python • React • Node.js • AWS • Docker"
        # Group into rows of ~6-8 per line
        chunk_size = 6
        for i in range(0, len(skills), chunk_size):
            chunk = skills[i:i+chunk_size]
            line = " • ".join(chunk)
            if i + chunk_size < len(skills):
                line += "\n"
            run = skills_para.add_run(line)
            run.font.name = FONT_NAME
            run.font.size = Pt(9.5)
            run.font.color.rgb = COLOR_BODY
    
    # ── PROFESSIONAL EXPERIENCE ───────────────────────────────────────────────
    has_experience = (
        experience or
        resume_data.get("experience_raw", "")
    )
    
    if has_experience:
        add_section_header(doc, "Professional Experience")
        
        if experience and isinstance(experience, list) and len(experience) > 0:
            # We have structured experience entries
            for entry in experience:
                company = entry.get("company", "")
                role = entry.get("role", "")
                dates = entry.get("dates", "")
                location = entry.get("location", "")
                bullets = entry.get("bullets", [])
                raw = entry.get("raw", "")
                
                # Company / Role header line
                if company or role or raw:
                    job_para = doc.add_paragraph()
                    set_paragraph_spacing(job_para, before=6, after=1)
                    
                    display_line = ""
                    if company and role:
                        display_line = f"{company} — {role}"
                    elif company:
                        display_line = company
                    elif role:
                        display_line = role
                    elif raw:
                        display_line = raw[:100]
                    
                    if display_line:
                        run = job_para.add_run(display_line)
                        run.bold = True
                        run.font.name = FONT_NAME
                        run.font.size = Pt(10.5)
                        run.font.color.rgb = COLOR_SUBHEADING
                    
                    if location:
                        run = job_para.add_run(f" | {location}")
                        run.font.name = FONT_NAME
                        run.font.size = Pt(10)
                        run.font.color.rgb = COLOR_META
                
                # Dates line
                if dates:
                    dates_para = doc.add_paragraph()
                    set_paragraph_spacing(dates_para, before=0, after=2)
                    run = dates_para.add_run(dates)
                    run.italic = True
                    run.font.name = FONT_NAME
                    run.font.size = Pt(9.5)
                    run.font.color.rgb = COLOR_META
                
                # Bullet points
                for bullet in bullets:
                    if bullet.strip():
                        bullet_para = doc.add_paragraph(style='List Bullet')
                        bullet_para.paragraph_format.left_indent = Inches(0.25)
                        bullet_para.paragraph_format.first_line_indent = Inches(-0.2)
                        set_paragraph_spacing(bullet_para, before=1, after=1)
                        
                        run = bullet_para.add_run(bullet.strip())
                        run.font.name = FONT_NAME
                        run.font.size = Pt(10)
                        run.font.color.rgb = COLOR_BODY
        
        else:
            # Fallback: render raw experience text
            exp_raw = resume_data.get("experience_raw", "")
            if exp_raw:
                for line in exp_raw.split('\n'):
                    line = line.strip()
                    if not line:
                        continue
                    
                    is_bullet = line.startswith(('•', '-', '–', '●', '*'))
                    is_date = bool(re.search(r'\d{4}', line))
                    
                    if is_bullet:
                        p = doc.add_paragraph(style='List Bullet')
                        p.paragraph_format.left_indent = Inches(0.25)
                        p.paragraph_format.first_line_indent = Inches(-0.2)
                        set_paragraph_spacing(p, before=1, after=1)
                        clean_line = re.sub(r'^[•\-–●\*·▪]\s*', '', line)
                        run = p.add_run(clean_line)
                        run.font.name = FONT_NAME
                        run.font.size = Pt(10)
                        run.font.color.rgb = COLOR_BODY
                    else:
                        p = doc.add_paragraph()
                        set_paragraph_spacing(p, before=4, after=1)
                        run = p.add_run(line)
                        run.bold = is_date
                        run.font.name = FONT_NAME
                        run.font.size = Pt(10 if not is_date else 10.5)
                        run.font.color.rgb = COLOR_SUBHEADING if not is_date else COLOR_SUBHEADING
    
    # ── EDUCATION ─────────────────────────────────────────────────────────────
    if education:
        add_section_header(doc, "Education")
        for line in education.split('\n'):
            line = line.strip()
            if line:
                p = doc.add_paragraph()
                set_paragraph_spacing(p, before=2, after=2)
                is_degree = any(kw in line.lower() for kw in
                                ['bachelor', 'master', 'phd', 'b.s.', 'm.s.', 'degree', 'diploma'])
                run = p.add_run(line)
                run.bold = is_degree
                run.font.name = FONT_NAME
                run.font.size = Pt(10)
                run.font.color.rgb = COLOR_SUBHEADING if is_degree else COLOR_BODY
    
    # ── PROJECTS ─────────────────────────────────────────────────────────────
    if projects:
        add_section_header(doc, "Key Projects")
        for line in projects.split('\n'):
            line = line.strip()
            if not line:
                continue
            is_bullet = line.startswith(('•', '-', '–', '●', '*'))
            if is_bullet:
                p = doc.add_paragraph(style='List Bullet')
                p.paragraph_format.left_indent = Inches(0.25)
                p.paragraph_format.first_line_indent = Inches(-0.2)
                set_paragraph_spacing(p, before=1, after=1)
                clean_line = re.sub(r'^[•\-–●\*·▪]\s*', '', line)
                run = p.add_run(clean_line)
            else:
                p = doc.add_paragraph()
                set_paragraph_spacing(p, before=4, after=1)
                run = p.add_run(line)
                run.bold = True
            run.font.name = FONT_NAME
            run.font.size = Pt(10)
            run.font.color.rgb = COLOR_BODY
    
    # ── CERTIFICATIONS ────────────────────────────────────────────────────────
    if certifications:
        add_section_header(doc, "Certifications & Credentials")
        for line in certifications.split('\n'):
            line = line.strip()
            if not line:
                continue
            is_bullet = line.startswith(('•', '-', '–'))
            if is_bullet:
                p = doc.add_paragraph(style='List Bullet')
                p.paragraph_format.left_indent = Inches(0.25)
                p.paragraph_format.first_line_indent = Inches(-0.2)
                set_paragraph_spacing(p, before=1, after=1)
                clean_line = re.sub(r'^[•\-–●\*·▪]\s*', '', line)
                run = p.add_run(clean_line)
            else:
                p = doc.add_paragraph()
                set_paragraph_spacing(p, before=2, after=2)
                run = p.add_run(line)
            run.font.name = FONT_NAME
            run.font.size = Pt(10)
            run.font.color.rgb = COLOR_BODY
    
    # ── Save to bytes ──────────────────────────────────────────────────────────
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.read()
